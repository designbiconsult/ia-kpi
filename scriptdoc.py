from docx import Document

doc = Document()
doc.add_heading('IA-KPI – Documentação de Arquitetura e Contexto do Projeto', 0)

doc.add_heading('1. Visão Geral do Projeto', level=1)
doc.add_paragraph(
    "O IA-KPI é um sistema web que conecta bases de dados empresariais (principalmente MySQL), sincroniza tabelas selecionadas "
    "e permite criar, visualizar e customizar indicadores de performance (KPIs) por setor (Financeiro, Comercial, Produção etc). "
    "A arquitetura é pensada para BI self-service, flexível e totalmente personalizável, com capacidade de relacionamento de "
    "tabelas e visualização amigável."
)

doc.add_heading('2. Stack e Estrutura', level=1)
doc.add_paragraph("""
- Backend: FastAPI (Python), SQLite para dados internos de controle, pymysql para conexão MySQL remota.
- Frontend: ReactJS, componentes separados (Login, Cadastro, Dashboard, Sincronismo, Configuração etc).
- Sincronismo: As tabelas do MySQL do cliente são lidas e copiadas para SQLite local apenas quando o usuário decide.
- Relacionamento: Modelagem visual inspirada no Power BI, permitindo relacionamentos 1:1, 1:N, N:N via drag and drop.
- KPIs: Painéis por setor, configuráveis, alimentados a partir das tabelas sincronizadas.
- Segurança: Cada usuário só acessa sua própria configuração e base.
- Arquivos de código: Separação clara backend/frontend, convenções modernas, documentação inline.
""")

doc.add_heading('3. Fluxo Principal do Usuário', level=1)
doc.add_paragraph("1. Cadastro/Login\n   - Usuário cria conta (salva no SQLite).\n   - Faz login, cai no Dashboard de indicadores (Financeiro, Comercial, Produção).")
doc.add_paragraph("2. Configuração da Conexão\n   - Formulário salva dados de conexão MySQL.\n   - Dados persistem para reuso e autenticação.")
doc.add_paragraph("3. Sincronismo de Tabelas/Views\n   - Usuário visualiza e seleciona tabelas/views do banco MySQL (trazidas via pymysql).\n   - Sincronismo executa cópia de estrutura/dados para SQLite local.\n   - Apenas tabelas/views marcadas ficam disponíveis para criação de indicadores e relacionamento.")
doc.add_paragraph("4. Relacionamento Visual\n   - Interface drag-and-drop para criar/gerenciar relacionamentos entre tabelas sincronizadas.\n   - Mapeamentos persistidos em tabela `relacionamentos` (com tipo 1:1, 1:N, etc).")
doc.add_paragraph("5. KPIs Básicos e Avançados\n   - Indicadores prontos por setor, usando SQL gerado automaticamente sobre a base SQLite sincronizada.\n   - Possibilidade de customizar e mapear indicadores específicos do negócio.")

doc.add_heading('4. Pontos Críticos/Diferenciais', level=1)
doc.add_paragraph("""
- Sem dependência de Streamlit: Sistema 100% web, independente de notebooks.
- Sincronismo seguro: Nunca consulta o MySQL do cliente em tempo real; apenas importa tabelas/views autorizadas.
- Relacionamento visual intuitivo: Usuário/admin pode ajustar relacionamentos como no Power BI.
- Personalização: Setores, KPIs, mapeamentos – tudo customizável.
""")

doc.add_heading('5. Comandos de Execução', level=1)
doc.add_paragraph("Backend:")
doc.add_paragraph("cd backend\npip install -r requirements.txt\npython -m uvicorn main:app --reload", style='Intense Quote')
doc.add_paragraph("Frontend:")
doc.add_paragraph("cd frontend\nnpm install\nnpm start", style='Intense Quote')

doc.add_heading('6. Observações Técnicas', level=1)
doc.add_paragraph("""
- Arquivo de conexão: database.db (em backend)
- Usuário SQLite: cada usuário tem campos próprios para salvar suas credenciais de banco remoto
- Persistência: Todas configurações e logs de relacionamento, sincronismo e indicadores salvos no SQLite
- Sincronização: Testada usando pymysql, suporte a views e tables do schema
- Validação: Validação robusta dos dados de conexão e checagem da existência do usuário
""")

doc.add_heading('7. Próximos Passos', level=1)
doc.add_paragraph("""
- Melhorar interface de relacionamento visual
- Habilitar criação dinâmica de KPIs customizados por arrasto
- Logs de sincronismo e validação de estrutura dinâmica
- Expor documentação OpenAPI sempre em /docs no backend
""")

doc.add_paragraph("Se precisar reiniciar o projeto ou abrir um novo chat:")
doc.add_paragraph("1. Cole esse conteúdo no início da conversa para garantir contexto total.\n2. Ou salve como .docx/.md e suba o arquivo no novo chat.")

doc.save("IA-KPI-Arquitetura-2025.docx")
print("Arquivo 'IA-KPI-Arquitetura-2025.docx' gerado com sucesso!")
